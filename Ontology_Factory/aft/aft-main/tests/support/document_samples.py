from __future__ import annotations

from io import BytesIO

from docx import Document

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_CONTENT_TYPE = "application/pdf"


def build_docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Overview", level=1)
    document.add_paragraph("Payment orchestration validates invoice references before settlement.")
    document.add_heading("Rules", level=2)
    document.add_paragraph("Every payment requires a signed approval token before execution.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Status"
    table.rows[1].cells[1].text = "Approved"

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_pdf_bytes(text: str) -> bytes:
    escaped_text = _escape_pdf_text(text.replace("\n", " "))
    content_stream = "\n".join(
        [
            "BT",
            "/F1 12 Tf",
            "72 720 Td",
            f"({escaped_text}) Tj",
            "ET",
        ]
    ).encode("ascii")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content_stream)).encode("ascii") + b" >>\nstream\n" + content_stream + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    chunks = [b"%PDF-1.4\n"]
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(sum(len(chunk) for chunk in chunks))
        chunks.append(f"{index} 0 obj\n".encode("ascii"))
        chunks.append(obj)
        chunks.append(b"\nendobj\n")

    body = b"".join(chunks)
    xref_offset = len(body)
    xref_entries = [b"0000000000 65535 f \n"]
    xref_entries.extend(f"{offset:010d} 00000 n \n".encode("ascii") for offset in offsets[1:])
    xref = b"xref\n0 6\n" + b"".join(xref_entries)
    trailer = f"trailer\n<< /Root 1 0 R /Size 6 >>\nstartxref\n{xref_offset}\n%%EOF".encode("ascii")
    return body + xref + trailer


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
