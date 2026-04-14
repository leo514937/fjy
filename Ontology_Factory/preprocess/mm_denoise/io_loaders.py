from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, Optional

from bs4 import BeautifulSoup
from charset_normalizer import from_bytes


@dataclass(frozen=True)
class LoadedDocument:
    path: Path
    text: str


def _decode_bytes(b: bytes, encoding_fallbacks: Iterable[str]) -> str:
    # 1) user-provided fallbacks
    for enc in encoding_fallbacks:
        try:
            return b.decode(enc)
        except Exception:
            pass
    # 2) best effort detection
    try:
        best = from_bytes(b).best()
        if best is not None:
            return str(best)
    except Exception:
        pass
    # 3) last resort
    return b.decode("utf-8", errors="replace")


def load_document(path: str, encoding_fallbacks: Iterable[str]) -> LoadedDocument:
    p = Path(path)
    if not p.exists():
        raise FileNotFoundError(str(p))

    suffix = p.suffix.lower()
    if suffix in {".txt", ".md"}:
        b = p.read_bytes()
        text = _decode_bytes(b, encoding_fallbacks)
        return LoadedDocument(path=p, text=text)

    if suffix in {".html", ".htm"}:
        b = p.read_bytes()
        html = _decode_bytes(b, encoding_fallbacks)
        soup = BeautifulSoup(html, "lxml")
        # Remove script/style/nav-like noisy blocks
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text("\n")
        return LoadedDocument(path=p, text=text)

    if suffix == ".docx":
        from docx import Document

        doc = Document(str(p))
        parts = []
        for para in doc.paragraphs:
            parts.append(para.text)
        text = "\n".join(parts)
        return LoadedDocument(path=p, text=text)

    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(p))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                parts.append("")
        text = "\n".join(parts)
        return LoadedDocument(path=p, text=text)

    raise ValueError(f"Unsupported file type: {suffix}")


def discover_inputs(input_globs: Iterable[str], base_dir: str) -> list[str]:
    base = Path(base_dir)
    paths: set[Path] = set()
    for g in input_globs:
        for p in base.glob(g):
            if p.is_file():
                paths.add(p.resolve())
    return [str(p) for p in sorted(paths)]


_CONTROL_CHARS = re.compile(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F]")


def normalize_text_for_pipeline(text: str) -> str:
    # Keep it minimal and reversible-ish.
    t = text.replace("\r\n", "\n").replace("\r", "\n")
    t = _CONTROL_CHARS.sub("", t)
    # Normalize full-width spaces
    t = t.replace("\u3000", " ")
    return t

